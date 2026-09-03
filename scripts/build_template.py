"""
Builds templates/eviction-petition.docx (a docxtemplater template) from the
lawyer's original file scripts/source-eviction-petition.docx.

Every yellow-highlighted span in the original is a variable. This script
replaces those spans with {placeholders}, removes the sample plaintiff list,
wraps the two optional demands in paragraph-level conditions, and strips all
highlighting so the generated document is clean.

Run:  python3 scripts/build_template.py
"""

import copy
import os

from docx import Document

SRC = os.path.join(os.path.dirname(__file__), "source-eviction-petition.docx")
OUT = os.path.join(
    os.path.dirname(__file__), "..", "templates", "eviction-petition.docx"
)


def set_span(paragraph, start, end, text):
    """Put `text` in run `start`, blank runs start+1..end (inclusive)."""
    runs = paragraph.runs
    runs[start].text = text
    for i in range(start + 1, end + 1):
        runs[i].text = ""


def strip_highlights(doc):
    for p in iter_all_paragraphs(doc):
        for r in p.runs:
            r.font.highlight_color = None


def iter_all_paragraphs(doc):
    for p in doc.paragraphs:
        yield p
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p
    for s in doc.sections:
        for p in list(s.header.paragraphs) + list(s.footer.paragraphs):
            yield p


def delete_paragraph(paragraph):
    el = paragraph._element
    el.getparent().remove(el)


def tag_paragraph_like(model, tag):
    """A paragraph containing only `tag`, cloned from `model` for formatting."""
    new_p = copy.deepcopy(model._element)
    for child in new_p.findall(
        "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}r"
    ):
        new_p.remove(child)
    model._element.addprevious(new_p)
    from docx.text.paragraph import Paragraph

    p = Paragraph(new_p, model._parent)
    run = p.add_run(tag)
    run.font.highlight_color = None
    return p


def main():
    doc = Document(SRC)

    body = doc.paragraphs
    cell = doc.tables[0].rows[0].cells[0]

    # ---------- header table: plaintiffs ----------
    # "بناء على طلب" stays; the highlighted remainder becomes one field.
    set_span(cell.paragraphs[1], 2, 6, "/ {plaintiff_name}")

    # the five sample heirs (List Paragraph items) are sample data, not template
    for p in list(cell.paragraphs[2:7]):
        delete_paragraph(p)

    # ---------- header table: defendant ----------
    cp = cell.paragraphs
    # paragraph indices shifted by 5 after the deletions above
    set_span(cp[4], 0, 11, "{defendant_name}")  # was p9
    set_span(cp[5], 0, 22, "ويعلن في: {defendant_address}")  # was p10

    # ---------- body: lease, premises, rent ----------
    set_span(body[3], 1, 8, "{lease_date_phrase}")
    set_span(body[3], 10, 23, "{premises_lead} {premises_address}")
    set_span(
        body[3],
        26,
        32,
        "( {property_use} ) لقاء أجرة شهرية قدرها {monthly_rent} د.ك "
        "(فقط {monthly_rent_words} شهريا)",
    )

    # ---------- body: date non-payment began ----------
    set_span(
        body[4],
        1,
        5,
        "امتنع المعلن إليه عن سداد الأجرة الشهرية من تاريخ {nonpayment_start_date}",
    )

    # ---------- body: arrears calculation ----------
    set_span(
        body[5],
        0,
        16,
        "وعليه فإن المعلن إليه لم يقم بسداد الأجرة عن الأشهر من "
        "{arrears_from_month} إلى {arrears_to_month} /{arrears_year} "
        "({monthly_rent} × {arrears_months_count} شهر = {arrears_total} د.ك "
        "{arrears_total_words})",
    )

    # ---------- demands ----------
    set_span(
        body[12],
        0,
        7,
        "{ordinal_eviction}: بإلزام المعلن إليه بإخلاء العين المبينة الوصف في صدر "
        "هذه الصحيفة وعقد الإيجار المؤرخ في {lease_date}.",
    )
    set_span(
        body[13],
        0,
        22,
        "{ordinal_arrears}: بأداء الأجرة المتأخرة للطالب عن الأشهر "
        "{arrears_months_list} من سنة {arrears_year} بمبلغ {monthly_rent} د.ك "
        "لكل شهر بما مجموعه {arrears_total} د.ك ({arrears_total_words}) عن الأشهر "
        "الفائتة، بالإضافة إلى ما يستجد من مقابل الانتفاع من تاريخ "
        "{benefit_start_date} حتى تمام الإخلاء.",
    )
    set_span(
        body[14],
        0,
        0,
        "{ordinal_penalty}: إلزام المعلن إليه بأداء مبلغ وقدره {penalty_amount} د.ك. "
        "({penalty_amount_words}) للطالب إعمالا للشرط الجزائي المقرر في البند السادس "
        "من عقد الإيجار المبرم بينهما.",
    )

    # رابعا -> computed ordinal (this demand is always present)
    for r in body[15].runs:
        if "رابعا" in r.text:
            r.text = r.text.replace("رابعا", "{ordinal_costs}", 1)
            break

    # ---------- optional demands become paragraph-level conditions ----------
    # (docxtemplater paragraphLoop: a paragraph holding only a tag is removed)
    tag_paragraph_like(body[12], "{#include_eviction_request}")
    tag_paragraph_like(body[13], "{/include_eviction_request}")
    tag_paragraph_like(body[14], "{#include_penalty_clause}")
    tag_paragraph_like(body[15], "{/include_penalty_clause}")

    strip_highlights(doc)

    os.makedirs(os.path.dirname(OUT), exist_ok=True)
    doc.save(OUT)
    print("wrote", os.path.normpath(OUT))


if __name__ == "__main__":
    main()
